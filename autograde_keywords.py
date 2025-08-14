import json
from docx import Document

# Load questions and keywords
with open('questions.json', 'r') as f:
    questions = json.load(f)

# Load student answers (example: answers.json)
with open('answers.json', 'r') as f:
    answers = json.load(f)

results = []
doc = Document()
doc.add_heading('Assignment Results', 0)

for q in questions:
    qid = str(q['id'])
    question = q['question']
    keywords = q['keywords']
    answer = answers.get(qid, "")
    matched = [kw for kw in keywords if kw.lower() in answer.lower()]
    score = len(matched) / len(keywords) if keywords else 0
    results.append({
        'question': question,
        'answer': answer,
        'matched_keywords': matched,
        'score': score
    })
    doc.add_heading(f"Q{qid}: {question}", level=1)
    doc.add_paragraph(f"Answer: {answer}")
    doc.add_paragraph(f"Matched keywords: {', '.join(matched)}")
    doc.add_paragraph(f"Score: {score:.2f}")

doc.save('graded_assignment.docx')

# Print results for manual review
for r in results:
    print(f"Question: {r['question']}")
    print(f"Answer: {r['answer']}")
    print(f"Matched keywords: {r['matched_keywords']}")
    print(f"Score: {r['score']:.2f}")
    print()
