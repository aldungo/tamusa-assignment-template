#!/usr/bin/env python3
"""
CSCI 1436 Programming Fundamentals I - Assignment #1
Student Submission Script

This script will guide you through answering assignment questions
and prepare your submission for both GitHub Classroom and Turnitin.
"""

from docx import Document
from docx.shared import Pt
import json
import datetime

# Assignment questions
questions = [
    {
        "id": "q1_definitions",
        "points": 40,
        "question": "1. (40 points) Briefly define (one to two sentences) each of the following ten terms:\n- REPL\n- Java API\n- JDK\n- Variable\n- Data type\n- Variable declaration\n- Assignment statement\n- Expression\n- Constant\n- Literal"
    },
    {
        "id": "q2_high_level", 
        "points": 12,
        "question": "2. (12 points) Why is it ideal to write programs in a high-level language instead of a low-level language?"
    },
    {
        "id": "q3_declarations",
        "points": 12,
        "question": "3. (12 points) Provide variable or constant declarations and initializations for:\n- double length = 23.6\n- float width = 14.7\n- long distance = 3172900000\n- final int speedOfSound = 343"
    },
    {
        "id": "q4_modulo",
        "points": 12,
        "question": "4. (12 points) Indicate the result of performing the remainder operation:\n- 40 % 17\n- 120 % 60\n- (65 + 8) % 25\n- (97 * 3 + 6) % 25"
    },
    {
        "id": "q5_sphere_volume",
        "points": 12,
        "question": "5. (12 points) Write a short program that requests a radius and calculates the volume of a sphere using: (4/3) * π * r^3"
    },
    {
        "id": "q6_debugging",
        "points": 12,
        "question": "6. (12 points) Identify the problem in each code snippet:\n- int userInput = 20;\n- double calculation = 50 / (userInput – 20);\n- int total = 2000000000;\n- int withdrawal = 1000000000;\n- total += withdrawal;\n- long toPluto = 3315000000L;\n- int toTheSun = 92955807;\n- int totalDistance = toPluto + toTheSun;"
    }
]

def get_student_info():
    """Get student information"""
    print("=== CSCI 1436 Programming Fundamentals I - Assignment #1 ===")
    print("Please enter your information:")
    print("=" * 60)
    
    first_name = input("Enter your first name: ").strip()
    last_name = input("Enter your last name: ").strip()
    student_id = input("Enter your student ID: ").strip()
    
    return first_name, last_name, student_id

def collect_responses():
    """Collect student responses"""
    responses = {}
    
    print(f"\nPlease answer all {len(questions)} questions:")
    print("For code questions, you can write multiple lines. Type 'END' on a new line when finished.")
    print("=" * 60)
    
    for i, question_data in enumerate(questions, 1):
        print(f"\nQuestion {i} ({question_data['points']} points):")
        print(question_data['question'])
        print("-" * 50)
        
        # Check if this might be a code question
        if 'program' in question_data['question'].lower() or 'code' in question_data['question'].lower():
            print("💡 Tip: For code answers, type your code and then 'END' on a new line when finished.")
            response_lines = []
            while True:
                line = input()
                if line.strip() == 'END':
                    break
                response_lines.append(line)
            response = '\n'.join(response_lines)
        else:
            response = input("Your answer: ")
        
        responses[question_data['id']] = {
            'question_num': i,
            'response': response,
            'timestamp': datetime.datetime.now().isoformat()
        }
        
        print("✅ Response recorded!")
    
    return responses

def create_word_document(first_name, last_name, student_id, responses):
    """Create Word document for Turnitin submission"""
    doc = Document()
    
    # Header
    doc.add_heading("CSCI 1436 Programming Fundamentals I", level=1)
    doc.add_heading("Assignment #1", level=2)
    doc.add_paragraph(f"Student: {first_name} {last_name}")
    doc.add_paragraph(f"Student ID: {student_id}")
    doc.add_paragraph(f"Submission Date: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph("")
    
    # Questions and responses
    for question_data in questions:
        response_data = responses[question_data['id']]
        
        doc.add_paragraph(f"Question {response_data['question_num']}:", style='Heading 3')
        doc.add_paragraph(question_data['question'])
        
        # Add response
        response_para = doc.add_paragraph("Answer:")
        response_para.add_run(f"\n{response_data['response']}")
        response_para.runs[-1].font.size = Pt(11)
        doc.add_paragraph("")
    
    # Save document
    filename = f"Assignment1_{first_name}_{last_name}.docx"
    doc.save(filename)
    return filename

def create_submission_data(first_name, last_name, student_id, responses):
    """Create submission data for autograding"""
    submission_data = {
        "student_info": {
            "first_name": first_name,
            "last_name": last_name,
            "student_id": student_id
        },
        "assignment_info": {
            "course": "CSCI 1436",
            "assignment": "Assignment #1",
            "submission_timestamp": datetime.datetime.now().isoformat()
        },
        "responses": responses
    }
    
    # Save JSON file
    json_filename = f"submission_{first_name}_{last_name}.json"
    with open(json_filename, 'w') as f:
        json.dump(submission_data, f, indent=2)
    
    return json_filename

def main():
    """Main function"""
    try:
        # Get student information
        first_name, last_name, student_id = get_student_info()
        
        # Collect responses
        responses = collect_responses()
        
        # Create Word document for Turnitin
        word_file = create_word_document(first_name, last_name, student_id, responses)
        
        # Create submission data for autograding
        json_file = create_submission_data(first_name, last_name, student_id, responses)
        
        print(f"\n🎉 Assignment completed successfully!")
        print(f"📄 Word document created: {word_file}")
        print(f"📊 Submission data created: {json_file}")
        print(f"\n📚 Next steps:")
        print(f"   1. Commit and push your changes to GitHub:")
        print(f"      git add .")
        print(f"      git commit -m 'Complete Assignment #1'")
        print(f"      git push origin main")
        print(f"   2. Upload {word_file} to Turnitin")
        print(f"   3. Check the 'Actions' tab in GitHub to see your auto-grading results!")
        
    except KeyboardInterrupt:
        print("\n❌ Assignment submission cancelled.")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("Please contact your instructor if you continue to have issues.")

if __name__ == "__main__":
    main()
