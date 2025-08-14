from docx import Document
from docx.shared import Pt

# Define the assignment questions
questions = [
    "1. (40 points) Briefly define (one to two sentences) each of the following ten terms:\n- REPL\n- Java API\n- JDK\n- Variable\n- Data type\n- Variable declaration\n- Assignment statement\n- Expression\n- Constant\n- Literal",
    "2. (12 points) Why is it ideal to write programs in a high-level language instead of a low-level language?",
    "3. (12 points) Provide variable or constant declarations and initializations for:\n- double length = 23.6\n- float width = 14.7\n- long distance = 3172900000\n- final int speedOfSound = 343",
    "4. (12 points) Indicate the result of performing the remainder operation:\n- 40 % 17\n- 120 % 60\n- (65 + 8) % 25\n- (97 * 3 + 6) % 25",
    "5. (12 points) Write a short program that requests a radius and calculates the volume of a sphere using: (4/3) * π * r^3",
    "6. (12 points) Identify the problem in each code snippet:\n- int userInput = 20;\n- double calculation = 50 / (userInput – 20);\n- int total = 2000000000;\n- int withdrawal = 1000000000;\n- total += withdrawal;\n- long toPluto = 3315000000L;\n- int toTheSun = 92955807;\n- int totalDistance = toPluto + toTheSun;"
]

# Get student info
first_name = input("Enter your first name: ").strip()
last_name = input("Enter your last name: ").strip()

# Create the document
doc = Document()
doc.add_heading("CSCI 1436 Programming Fundamentals I", level=1)
doc.add_heading("Assignment #1", level=2)
doc.add_paragraph(f"Student: {first_name} {last_name}")
doc.add_paragraph("")

# Ask questions and collect responses
for i, question in enumerate(questions, start=1):
    doc.add_paragraph(f"Question {i}:", style='Heading 3')
    doc.add_paragraph(question)
    response = input(f"Your response to Question {i}: ")
    response_para = doc.add_paragraph()
    response_run = response_para.add_run(response)
    response_run.font.size = Pt(11)
    doc.add_paragraph("")

# Save the file
filename = f"Assignment1_{first_name}_{last_name}.docx"
doc.save(filename)
print(f"✅ Assignment saved as '{filename}'. You can now submit it to Blackboard or Turnitin.")
