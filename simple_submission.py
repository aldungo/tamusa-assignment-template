#!/usr/bin/env python3
"""
CSCI 1436 Programming Fundamentals I - Assignment #1
Simple Student Submission Script

Complete the assignment by answering questions.
Generates: JSON for autograding + Word doc for Blackboard submission proof.
"""

from docx import Document
from docx.shared import Pt
import json
import datetime

def main():
    print("=== CSCI 1436 Programming Fundamentals I - Assignment #1 ===")
    print("Complete all questions. Your answers will be auto-graded!")
    print("=" * 60)
    
    # Get student info
    first_name = input("Enter your first name: ").strip()
    last_name = input("Enter your last name: ").strip()
    student_id = input("Enter your student ID: ").strip()
    
    # Assignment questions
    questions = [
        {
            "id": "q1_definitions",
            "points": 40,
            "text": "1. (40 points) Briefly define (one to two sentences) each of the following ten terms:\n- REPL\n- Java API\n- JDK\n- Variable\n- Data type\n- Variable declaration\n- Assignment statement\n- Expression\n- Constant\n- Literal"
        },
        {
            "id": "q2_languages", 
            "points": 12,
            "text": "2. (12 points) Why is it ideal to write programs in a high-level language instead of a low-level language?"
        },
        {
            "id": "q3_declarations",
            "points": 12,
            "text": "3. (12 points) Provide variable or constant declarations and initializations for:\n- double length = 23.6\n- float width = 14.7\n- long distance = 3172900000\n- final int speedOfSound = 343"
        },
        {
            "id": "q4_modulo",
            "points": 12,
            "text": "4. (12 points) Indicate the result of performing the remainder operation:\n- 40 % 17\n- 120 % 60\n- (65 + 8) % 25\n- (97 * 3 + 6) % 25"
        },
        {
            "id": "q5_programming",
            "points": 12,
            "text": "5. (12 points) Write a short program that requests a radius and calculates the volume of a sphere using: (4/3) * π * r^3"
        },
        {
            "id": "q6_debugging",
            "points": 12,
            "text": "6. (12 points) Identify the problem in each code snippet:\n- int userInput = 20;\n- double calculation = 50 / (userInput – 20);\n- int total = 2000000000;\n- int withdrawal = 1000000000;\n- total += withdrawal;\n- long toPluto = 3315000000L;\n- int toTheSun = 92955807;\n- int totalDistance = toPluto + toTheSun;"
        }
    ]
    
    # Collect answers
    answers = {}
    total_points = sum(q["points"] for q in questions)
    
    print(f"\nAnswer all {len(questions)} questions (Total: {total_points} points):")
    print("=" * 50)
    
    for i, question in enumerate(questions, 1):
        print(f"\nQuestion {i} ({question['points']} points):")
        print(question["text"])
        print("-" * 40)
        
        # Multi-line input for code questions
        if "program" in question["text"].lower() or "code" in question["text"].lower():
            print("💡 For code, type multiple lines then 'END' on a new line:")
            lines = []
            while True:
                line = input()
                if line.strip() == "END":
                    break
                lines.append(line)
            answer = "\n".join(lines)
        else:
            answer = input("Your answer: ")
        
        answers[question["id"]] = {
            "question_number": i,
            "answer": answer,
            "max_points": question["points"],
            "timestamp": datetime.datetime.now().isoformat()
        }
        
        print("✅ Answer saved!")
    
    # Create submission data for autograding
    submission = {
        "student": {
            "name": f"{first_name} {last_name}",
            "first_name": first_name,
            "last_name": last_name,
            "id": student_id
        },
        "assignment": {
            "course": "CSCI 1436",
            "name": "Assignment #1",
            "total_points": total_points,
            "submitted_at": datetime.datetime.now().isoformat()
        },
        "answers": answers
    }
    
    # Save JSON for autograding
    json_file = f"assignment1_{first_name}_{last_name}.json"
    with open(json_file, 'w') as f:
        json.dump(submission, f, indent=2)
    
    # Create simple Word doc for Blackboard submission proof
    doc = Document()
    doc.add_heading("CSCI 1436 Programming Fundamentals I", level=1)
    doc.add_heading("Assignment #1 - Submission Proof", level=2)
    doc.add_paragraph(f"Student: {first_name} {last_name}")
    doc.add_paragraph(f"Student ID: {student_id}")
    doc.add_paragraph(f"Submitted: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph("")
    doc.add_paragraph("This document serves as proof of assignment submission.")
    doc.add_paragraph("Answers have been submitted for autograding via GitHub Classroom.")
    doc.add_paragraph(f"Total Questions: {len(questions)} | Total Points: {total_points}")
    
    word_file = f"Assignment1_{first_name}_{last_name}.docx"
    doc.save(word_file)
    
    # Success message
    print(f"\n" + "="*60)
    print(f"🎉 Assignment completed!")
    print("="*60)
    print(f"📊 Autograding file: {json_file}")
    print(f"📄 Blackboard proof: {word_file}")
    
    # Auto-save to Git immediately for safety
    print(f"\n�️  AUTO-SAVING YOUR WORK TO GIT...")
    try:
        import subprocess
        subprocess.run(["git", "add", "."], check=True, capture_output=True)
        subprocess.run(["git", "commit", "-m", f"Auto-save: Assignment 1 completed by {first_name} {last_name}"], check=True, capture_output=True)
        print("✅ Work automatically saved to Git (local backup)!")
        print("⚠️  IMPORTANT: You still need to PUSH to submit for grading!")
    except Exception as e:
        print("⚠️  Could not auto-save to Git. Please save manually!")
        print(f"   Error: {e}")
    
    print(f"\n🔴 CRITICAL - PUSH NOW TO SUBMIT:")
    print(f"   git push")
    print(f"\n📚 Then upload {word_file} to Blackboard")
    print(f"🎯 Check GitHub Actions tab for your auto-grading score!")
    print(f"\n⚠️  WARNING: If you close Codespace without 'git push', your work may be lost!")
    print("="*60)

if __name__ == "__main__":
    main()
